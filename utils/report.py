import io
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _bold_run_paragraph(doc, text, size=11):
    """'**강조**' 마크다운 표기를 굵게 처리해서 문단에 추가한다."""
    p = doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(size)
        if i % 2 == 1:
            run.bold = True
    return p


def _fill_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def generate_cooling_report_docx(data: dict) -> bytes:
    """냉각수 시스템 현장 서비스 리포트를 .docx 바이트로 생성한다."""
    doc = Document()

    title = doc.add_heading('냉각수 시스템 현장 서비스 리포트', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"생성일시: {data.get('generated_at', datetime.now().strftime('%Y-%m-%d %H:%M'))}").font.size = Pt(9)

    doc.add_heading('1. 운전 조건 및 핵심 지수', level=2)
    _fill_table(
        doc,
        ["항목", "값"],
        [
            ["수온 (℃)", f"{data.get('temp', 0):.1f}"],
            ["pH (목표)", f"{data.get('ph', 0):.2f}"],
            ["농축배수 (Cycles)", f"{data.get('coc', 0):.1f}"],
            ["LSI (Bulk)", f"{data.get('lsi', 0):.2f}"],
            ["LSI (Skin)", f"{data.get('lsi_skin', 0):.2f}"],
            ["RSI", f"{data.get('rsi', 0):.2f}"],
            ["PSI", f"{data.get('psi', 0):.2f}"],
            ["Larson-Skold (L-S)", f"{data.get('ls_idx', 0):.2f}"],
        ]
    )

    stress = data.get('stress')
    if stress:
        doc.add_heading('2. 통합 스트레스 지수', level=2)
        _bold_run_paragraph(doc, f"**종합 점수: {stress['score']:.0f} / 100 ({stress['band']})**")
        bd = stress['breakdown']
        doc.add_paragraph(f"기여도 — LSI {bd['LSI']:.0f} · RSI {bd['RSI']:.0f} · PSI {bd['PSI']:.0f} · L-S {bd['L-S']:.0f} (100점 만점 환산)").runs[0].font.size = Pt(9)
        doc.add_paragraph("※ LSI/RSI/PSI/L-S 4개 지수를 가중합산한 자체 참고 지표이며, 특정 제조사의 특허 알고리즘을 재현한 것이 아닙니다.").runs[0].font.size = Pt(8)

    coupon = data.get('coupon')
    doc.add_heading('3. 부식쿠폰 실측 결과', level=2)
    if coupon:
        _fill_table(
            doc,
            ["재질", "실측치 (mpy)", "등급"],
            [
                ["Mild Steel", f"{coupon['ms']['mpy']:.1f}", coupon['ms']['grade']],
                ["Copper", f"{coupon['cu']['mpy']:.2f}", coupon['cu']['grade']],
            ]
        )
        if data.get('coupon_comment'):
            doc.add_paragraph(data['coupon_comment']).runs[0].font.size = Pt(9)
    else:
        doc.add_paragraph("실측 데이터 미입력").runs[0].font.size = Pt(9)

    doc.add_heading('4. 약품 선정 및 투입량', level=2)
    chem_rows = data.get('chem_rows', [])
    if chem_rows:
        _fill_table(
            doc,
            ["구분", "제품명", "주성분", "투입농도 (ppm)", "일일 사용량 (kg/day)"],
            [[r['category'], r['product'], r['ingredient'], f"{r['dosage']:.1f}", f"{r['daily_kg']:.1f}"] for r in chem_rows]
        )
    else:
        doc.add_paragraph("선정된 약품 없음").runs[0].font.size = Pt(9)

    if data.get('rec_reason'):
        doc.add_heading('5. 약품 선정 사유', level=2)
        _bold_run_paragraph(doc, data['rec_reason'])

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Water Master Pro — 자동 생성 리포트 (현장 확인 후 최종 처방을 확정하십시오)").font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
